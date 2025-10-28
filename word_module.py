from docx import Document
from docx.shared import Inches
import os, datetime, config

def append_to_word(counselor: str, data: dict, image_path: str = None):
    """Append claim info to counselor's Word document (one doc per counselor)."""
    try:
        folder = config.WORD_DIR
        os.makedirs(folder, exist_ok=True)
        path = os.path.join(folder, f"{counselor}.docx")
        doc = Document(path) if os.path.exists(path) else Document()
        if not os.path.exists(path):
            doc.add_heading(f"{counselor} — Claims", level=1)
        _entry(doc, data, image_path)
        doc.save(path)
    except Exception as e:
        print(f"Word write error: {e}")

def _entry(doc, d, image_path):
    doc.add_paragraph("")
    doc.add_paragraph(f"Client: {d.get('Client','Unknown')}")
    doc.add_paragraph(f"Insurance: {d.get('Insurance','N/A')}")
    doc.add_paragraph(f"Date of Service: {d.get('Date','N/A')}")
    doc.add_paragraph(f"Client Responsibility: ${d.get('Client Responsibility','0.00')}")
    doc.add_paragraph(f"Insurance Payment: ${d.get('Insurance Payment','0.00')}")
    doc.add_paragraph(f"Added: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
    
    # Add screenshot if provided
    if image_path and os.path.exists(image_path):
        try:
            doc.add_paragraph("Screenshot:")
            doc.add_picture(image_path, width=Inches(6.0))
        except Exception as e:
            doc.add_paragraph(f"(Could not embed image: {e})")
